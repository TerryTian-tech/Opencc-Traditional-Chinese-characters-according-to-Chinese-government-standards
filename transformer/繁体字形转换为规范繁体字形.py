import os
import sys
import tempfile
import shutil
import codecs
import chardet
import zipfile
import xml.etree.ElementTree as ET
import win32com.client as win32
import msvcrt

from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QPushButton, QTextEdit, QFileDialog, QLabel, QProgressBar,
                             QMessageBox, QGroupBox, QComboBox, QCheckBox, QLineEdit,
                             QStyleFactory, QTabWidget, QListWidget, QSplitter, QMenuBar,
                             QMenu, QAction, QActionGroup, QRadioButton, QButtonGroup)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QIcon, QFont, QPixmap, QColor, QPalette

from docx import Document
from opencc import OpenCC
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import re


class ConversionWorker(QThread):
    """
    转换工作线程，避免UI阻塞
    """
    progress_updated = pyqtSignal(int, str)  # 进度信号
    conversion_finished = pyqtSignal(bool, str)  # 完成信号
    log_message = pyqtSignal(str)  # 日志消息信号
    
    def __init__(self, input_path, output_folder, conversion_type='t2gov', preserve_format=True, 
                 convert_footnotes=True):
        super().__init__()
        self.input_path = input_path
        self.output_folder = output_folder
        self.conversion_type = conversion_type
        self.preserve_format = preserve_format
        self.convert_footnotes = convert_footnotes
    
    def detect_encoding(self, file_path):
        """检测文件编码，特别处理中文ANSI编码"""
        self.log_message.emit(f"检测文件编码: {file_path}")
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            
        # 首先尝试chardet检测
        result = chardet.detect(raw_data)
        encoding = result['encoding']
        confidence = result['confidence']
        
        self.log_message.emit(f"chardet检测结果: {encoding} (置信度: {confidence})")
        
        # 特别处理GB18030编码
        # 如果检测到GB2312，但文件可能包含GB18030特有字符，优先尝试GB18030
        if encoding == 'GB2312' and confidence < 0.95:
            try:
                # 尝试用GB18030解码整个文件
                decoded = raw_data.decode('gb18030', errors='strict')
                # 检查是否包含GB18030特有的字符范围
                # GB18030扩展了GB2312，支持更多汉字和符号
                if any(ord(char) > 0x9FFF for char in decoded):  # 检查是否包含扩展汉字
                    self.log_message.emit("检测到GB18030扩展字符，使用GB18030编码")
                    return 'gb18030'
                else:
                    # 虽然没有扩展字符，但为了兼容性，仍使用GB18030
                    self.log_message.emit("使用GB18030编码以确保兼容性")
                    return 'gb18030'
            except UnicodeDecodeError:
                # 如果GB18030解码失败，回退到检测到的编码
                pass
        
        # 如果置信度低或者是常见误判情况，尝试中文编码
        if confidence < 0.7 or encoding in ['ISO-8859-1', 'Windows-1252', 'ascii']:
            # 尝试常见中文编码，优先尝试GB18030
            chinese_encodings = ['gb18030', 'gbk', 'gb2312', 'big5']
            for enc in chinese_encodings:
                try:
                    # 尝试解码前1000个字节
                    test_data = raw_data[:1000]
                    decoded = test_data.decode(enc, errors='strict')
                    # 如果包含中文字符，认为可能是正确的编码
                    if any('\u4e00' <= char <= '\u9fff' for char in decoded):
                        self.log_message.emit(f"检测到中文字符，使用编码: {enc}")
                        return enc
                except UnicodeDecodeError:
                    continue
        
        # 如果检测到UTF-8但置信度不高，尝试GB18030
        if encoding == 'utf-8' and confidence < 0.9:
            try:
                # 尝试用GB18030解码
                decoded = raw_data.decode('gb18030', errors='strict')
                # 检查是否包含中文字符
                if any('\u4e00' <= char <= '\u9fff' for char in decoded):
                    self.log_message.emit("检测到GB18030编码的中文字符，使用GB18030编码")
                    return 'gb18030'
            except UnicodeDecodeError:
                pass
        
        # 默认使用检测到的编码，如果是None则使用utf-8
        if not encoding:
            encoding = 'utf-8'
        
        # 如果是GB2312，优先使用GB18030以确保兼容性
        if encoding.lower() in ['gb2312', 'gbk']:
            self.log_message.emit(f"将{encoding}升级为GB18030以确保更好的兼容性")
            return 'gb18030'
        
        return encoding

    def safe_read_file(self, file_path, encoding):
        """安全读取文件，处理编码问题"""
        # 优先尝试GB18030，因为它兼容GB2312和GBK
        if encoding.lower() in ['gb2312', 'gbk']:
            try:
                with codecs.open(file_path, 'r', encoding='gb18030', errors='strict') as f:
                    return f.read()
            except UnicodeDecodeError as e:
                self.log_message.emit(f"GB18030严格模式读取失败: {e}，尝试原编码")
        
        try:
            with codecs.open(file_path, 'r', encoding=encoding, errors='strict') as f:
                return f.read()
        except UnicodeDecodeError:
            # 如果严格模式失败，尝试使用errors='ignore'
            self.log_message.emit(f"使用严格模式读取失败，尝试忽略错误字符")
            try:
                with codecs.open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()
                    # 检查读取的内容是否包含有效的中文字符
                    if any('\u4e00' <= char <= '\u9fff' for char in content):
                        return content
                    else:
                        # 如果没有中文字符，可能是编码错误，尝试GB18030
                        self.log_message.emit("读取内容不包含中文字符，尝试GB18030编码")
                        with codecs.open(file_path, 'r', encoding='gb18030', errors='ignore') as f2:
                            return f2.read()
            except Exception as e:
                self.log_message.emit(f"读取文件时发生错误: {e}")
                # 最后尝试使用GB18030
                try:
                    with codecs.open(file_path, 'r', encoding='gb18030', errors='ignore') as f:
                        return f.read()
                except Exception as e2:
                    self.log_message.emit(f"最终读取失败: {e2}")
                    return ""

    def convert_doc_to_docx(self, input_path, output_folder):
        """
        将DOC文件转换为DOCX文件
        :param input_path: 输入的DOC文件路径
        :param output_folder: 输出文件夹路径
        :return: 转换后的DOCX文件路径或False
        """
        try:
            if not os.path.exists(input_path):
                self.log_message.emit(f"错误：文件不存在 - {input_path}")
                return False
                
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)
                self.log_message.emit(f"创建输出目录: {output_folder}")
            
            self.log_message.emit(f"正在转换DOC文件: {os.path.basename(input_path)}")
            
            # 创建Word应用程序实例
            word = win32.gencache.EnsureDispatch('Word.Application')
            word.Visible = False  # 不显示Word界面
            
            try:
                # 打开DOC文件
                doc = word.Documents.Open(input_path)
                
                # 生成输出文件名
                filename = os.path.basename(input_path)
                docx_filename = os.path.splitext(filename)[0] + ".docx"
                output_path = os.path.join(output_folder, docx_filename)
                
                # 保存为DOCX格式
                doc.SaveAs(output_path, FileFormat=16)  # 16代表docx格式
                doc.Close()
                
                self.log_message.emit(f"已转换为DOCX文件: {output_path}")
                return output_path
                
            except Exception as e:
                self.log_message.emit(f"转换DOC文件时出错: {str(e)}")
                return False
                
            finally:
                # 确保Word应用程序被关闭
                word.Quit()
                
        except Exception as e:
            self.log_message.emit(f"处理DOC文件 {input_path} 时出错: {str(e)}")
            return False

    def convert_txt_file(self, input_path, output_folder):
        """
        将txt文件转换为简体
        :param input_path: 输入文件路径
        :param output_folder: 输出文件夹路径
        :return: 转换后的文件路径或False
        """
        cc = OpenCC(self.conversion_type)
        
        try:
            if not os.path.exists(input_path):
                self.log_message.emit(f"错误：文件不存在 - {input_path}")
                return False
                
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)
                self.log_message.emit(f"创建输出目录: {output_folder}")
            
            self.log_message.emit(f"正在处理txt文件: {os.path.basename(input_path)}")
            
            # 检测文件编码
            encoding = self.detect_encoding(input_path)
            self.log_message.emit(f"最终使用的编码: {encoding}")
            
            # 读取文件内容
            content = self.safe_read_file(input_path, encoding)
            
            # 繁简转换
            converted_content = cc.convert(content)
            
            # 保存文件
            output_filename = f"convert_{os.path.basename(input_path)}"
            output_path = os.path.join(output_folder, output_filename)
            
            with codecs.open(output_path, 'w', encoding='utf-8') as f:
                f.write(converted_content)
            
            self.log_message.emit(f"已保存: {output_path}")
            return output_path
            
        except Exception as e:
            self.log_message.emit(f"处理txt文件 {input_path} 时出错: {str(e)}")
            return False

    class DocxTraditionalSimplifiedConverter:
        def __init__(self, worker, config='t2gov'):
            """
            初始化转换器
            - 't2gov': 繁体转规范繁体
            - 't2new': 繁体旧字形转新字形，但保留异体字不转换
            - 't2gov_keep_simp': 繁体转规范繁体，但保留文档内原有简体字
            - 't2new_keep_simp': 繁体旧字形转新字形，但保留文档内原有简体字和异体字
            - 't2s': 繁体转简体
            """
            self.worker = worker
            self.cc = OpenCC(config)
            self.config = config
        
        def convert_text(self, text):
            """转换文本内容"""
            if text and isinstance(text, str):
                return self.cc.convert(text)
            return text
        
        def _convert_footnotes_using_zip_manipulation(self, input_path, output_path):
            """
            通过直接操作docx文件（zip格式）来转换脚注和尾注
            这是最可靠的方法，因为它直接修改XML文件
            """
            try:
                # 创建临时目录
                with tempfile.TemporaryDirectory() as temp_dir:
                    # 解压docx文件
                    with zipfile.ZipFile(input_path, 'r') as zip_ref:
                        zip_ref.extractall(temp_dir)
                    
                    # 转换脚注XML文件
                    footnotes_path = os.path.join(temp_dir, 'word', 'footnotes.xml')
                    if os.path.exists(footnotes_path):
                        self._convert_xml_file(footnotes_path)
                        self.worker.log_message.emit("已转换脚注内容")
                    else:
                        self.worker.log_message.emit("文档中没有脚注")
                    
                    # 转换尾注XML文件（如果有）
                    endnotes_path = os.path.join(temp_dir, 'word', 'endnotes.xml')
                    if os.path.exists(endnotes_path):
                        self._convert_xml_file(endnotes_path)
                        self.worker.log_message.emit("已转换尾注内容")
                    
                    # 重新压缩为docx文件
                    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for root, dirs, files in os.walk(temp_dir):
                            for file in files:
                                file_path = os.path.join(root, file)
                                # 计算在zip中的相对路径
                                arcname = os.path.relpath(file_path, temp_dir)
                                zipf.write(file_path, arcname)
                    
                    return True
                    
            except Exception as e:
                self.worker.log_message.emit(f"通过zip操作转换脚注时出错: {e}")
                return False
        
        def _convert_xml_file(self, xml_path):
            """转换XML文件中的文本内容"""
            try:
                # 读取XML文件
                tree = ET.parse(xml_path)
                root = tree.getroot()
                
                # 定义XML命名空间
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
                }
                
                # 注册命名空间以便XPath查询
                for prefix, uri in namespaces.items():
                    ET.register_namespace(prefix, uri)
                
                # 查找所有文本节点
                text_elements = root.findall('.//w:t', namespaces)
                for elem in text_elements:
                    if elem.text:
                        elem.text = self.convert_text(elem.text)
                
                # 保存修改后的XML
                tree.write(xml_path, encoding='utf-8', xml_declaration=True)
                
            except Exception as e:
                self.worker.log_message.emit(f"转换XML文件 {xml_path} 时出错: {e}")
                # 如果XML解析失败，尝试使用正则表达式方法
                self._convert_xml_file_with_regex(xml_path)
        
        def _convert_xml_file_with_regex(self, xml_path):
            """使用正则表达式转换XML文件中的文本内容（备用方法）"""
            try:
                with open(xml_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # 使用正则表达式找到XML标签外的文本内容并转换
                def convert_text_in_xml(match):
                    # 匹配文本内容但不匹配标签属性
                    text = match.group(1)
                    # 只在文本包含中文字符时转换
                    if any('\u4e00' <= char <= '\u9fff' for char in text):
                        return '>' + self.convert_text(text) + '<'
                    else:
                        return match.group(0)
                
                # 匹配标签之间的文本内容
                pattern = r'>([^<]+?)<'
                converted_content = re.sub(pattern, convert_text_in_xml, content)
                
                with open(xml_path, 'w', encoding='utf-8') as f:
                    f.write(converted_content)
                    
            except Exception as e:
                self.worker.log_message.emit(f"使用正则表达式转换XML文件 {xml_path} 时出错: {e}")
        
        def convert_document(self, input_path, output_path=None):
            """
            转换整个Word文档，保留所有格式（包含脚注和尾注转换）
            """
            if output_path is None:
                filename, ext = os.path.splitext(input_path)
                output_path = f"convert_{filename}{ext}"
            
            self.worker.log_message.emit(f"开始转换文档: {input_path}")
            
            # 首先处理脚注和尾注（通过zip操作）
            temp_output = output_path + ".temp.docx"
            footnote_success = self._convert_footnotes_using_zip_manipulation(input_path, temp_output)
            
            if footnote_success:
                # 如果脚注转换成功，使用temp文件继续处理其他内容
                processing_file = temp_output
            else:
                # 如果脚注转换失败，使用原始文件
                self.worker.log_message.emit("脚注转换失败，将只转换正文内容")
                processing_file = input_path
                temp_output = None
            
            try:
                # 读取文档并转换其他内容
                doc = Document(processing_file)
                
                # 转换正文段落
                self._convert_paragraphs(doc.paragraphs)
                
                # 转换表格内容
                self._convert_tables(doc.tables)
                
                # 转换页眉
                for section in doc.sections:
                    self._convert_paragraphs(section.header.paragraphs)
                    # 转换页眉中的表格
                    if hasattr(section.header, 'tables') and section.header.tables:
                        self._convert_tables(section.header.tables)
                
                # 转换页脚
                for section in doc.sections:
                    self._convert_paragraphs(section.footer.paragraphs)
                    # 转换页脚中的表格
                    if hasattr(section.footer, 'tables') and section.footer.tables:
                        self._convert_tables(section.footer.tables)
                
                # 保存最终文档
                doc.save(output_path)
                self.worker.log_message.emit(f"转换完成，保存至: {output_path}")
                
                # 清理临时文件
                if temp_output and os.path.exists(temp_output):
                    os.remove(temp_output)
                
                return output_path
                
            except Exception as e:
                self.worker.log_message.emit(f"处理文档时出错: {e}")
                # 如果出错，尝试直接复制文件
                if temp_output and os.path.exists(temp_output):
                    shutil.copy2(temp_output, output_path)
                    if os.path.exists(temp_output):
                        os.remove(temp_output)
                    self.worker.log_message.emit(f"已保存基本转换的文档: {output_path}")
                    return output_path
                else:
                    # 最后的手段：直接复制原始文件
                    shutil.copy2(input_path, output_path)
                    self.worker.log_message.emit(f"转换失败，已复制原始文件到: {output_path}")
                    return output_path
        
        def _convert_paragraphs(self, paragraphs):
            """转换段落集合，保留所有格式"""
            for paragraph in paragraphs:
                self._convert_paragraph(paragraph)
        
        def _convert_paragraph(self, paragraph):
            """转换单个段落，保留所有run的格式"""
            if not paragraph.text.strip():
                return
            
            # 逐个处理run，保留每个run的独立格式
            for run in paragraph.runs:
                if run.text.strip():
                    original_text = run.text
                    converted_text = self.convert_text(original_text)
                    
                    # 保留原有格式的情况下更新文本
                    self._preserve_run_format(run, converted_text)
        
        def _preserve_run_format(self, run, new_text):
            """
            保留run的所有原始格式，只更新文本内容
            包括字体、大小、颜色、粗体、斜体、下划线等
            """
            # 保存当前格式
            original_bold = run.bold
            original_italic = run.italic
            original_underline = run.underline
            original_color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
            
            # 安全地获取高亮颜色
            original_highlight = None
            try:
                original_highlight = run.font.highlight_color
            except:
                pass
            
            # 保存字体信息
            original_font_name = run.font.name
            original_size = run.font.size
            
            # 更新文本内容
            run.text = new_text
            
            # 恢复格式
            run.bold = original_bold
            run.italic = original_italic
            run.underline = original_underline
            
            if original_color:
                if run.font.color is None:
                    run.font.color.rgb = original_color
                else:
                    run.font.color.rgb = original_color
            
            if original_highlight:
                try:
                    run.font.highlight_color = original_highlight
                except:
                    pass
            
            if original_font_name:
                run.font.name = original_font_name
                # 设置中文字体
                try:
                    if hasattr(run, '_element') and hasattr(run._element, 'rPr'):
                        rpr = run._element.rPr
                        if rpr is not None:
                            # 创建或获取字体设置
                            fonts = rpr.find(qn('w:rFonts'))
                            if fonts is None:
                                fonts = OxmlElement('w:rFonts')
                                rpr.append(fonts)
                            fonts.set(qn('w:eastAsia'), original_font_name)
                except Exception as e:
                    self.worker.log_message.emit(f"设置中文字体时出错: {e}")
            
            if original_size:
                run.font.size = original_size
        
        def _convert_tables(self, tables):
            """转换表格内容，保留表格格式"""
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        # 转换单元格中的段落
                        self._convert_paragraphs(cell.paragraphs)
                        
                        # 递归处理嵌套表格
                        for nested_table in cell.tables:
                            self._convert_tables([nested_table])

    def convert_docx_file(self, input_path, output_folder):
        """
        将Word文档转换为简体/繁体
        :param input_path: 输入文件路径
        :param output_folder: 输出文件夹路径
        :return: 转换后的文件路径
        """
        try:
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)
                self.log_message.emit(f"创建输出目录: {output_folder}")
            
            if os.path.isfile(input_path) and input_path.lower().endswith('.docx'):
                self.log_message.emit(f"正在处理: {os.path.basename(input_path)}")
                
                # 使用新的转换器类
                converter = self.DocxTraditionalSimplifiedConverter(self, self.conversion_type)
                output_path = os.path.join(output_folder, f"convert_{os.path.basename(input_path)}")
                converter.convert_document(input_path, output_path)
                
                self.log_message.emit(f"已保存: {output_path}")
                return output_path
                
            else:
                self.log_message.emit("错误：输入的路径不是有效的.docx文件")
                return False
                
        except Exception as e:
            self.log_message.emit(f"处理 {input_path} 时出错: {str(e)}")
            return False

    def process_files(self):
        """处理文件的主要逻辑"""
        self.progress_updated.emit(0, "开始处理...")
        
        # 处理单个文件
        if os.path.isfile(self.input_path):
            file_ext = os.path.splitext(self.input_path)[1].lower()
            
            if file_ext == '.docx':
                result = self.convert_docx_file(self.input_path, self.output_folder)
                if result:
                    self.progress_updated.emit(100, "转换完成!")
                    return True
                else:
                    return False
            elif file_ext == '.doc':
                # 先转换为DOCX，然后再进行繁简转换   
                self.log_message.emit("检测到DOC文件，先转换为DOCX格式...")
                
                # 创建临时目录用于存放临时转换的DOCX文件
                with tempfile.TemporaryDirectory() as temp_dir:
                    docx_path = self.convert_doc_to_docx(self.input_path, temp_dir)
                    if docx_path:
                        self.log_message.emit("DOC文件转换成功，开始繁简转换...")
                        result = self.convert_docx_file(docx_path, self.output_folder)
                        if result:
                            self.progress_updated.emit(100, "转换完成!")
                            return True
                        else:
                            return False
                    else:
                        self.log_message.emit("DOC文件转换失败")
                        return False
            elif file_ext == '.txt':
                result = self.convert_txt_file(self.input_path, self.output_folder)
                if result:
                    self.progress_updated.emit(100, "转换完成!")
                    return True
                else:
                    return False
            else:
                self.log_message.emit("错误：不支持的文件格式，仅支持docx、doc和txt文件")
                return False
        
        # 处理文件夹
        elif os.path.isdir(self.input_path):
            # 获取所有支持的文件
            supported_files = []
            for f in os.listdir(self.input_path):
                file_ext = os.path.splitext(f)[1].lower()
                if file_ext in ['.docx', '.doc', '.txt']:
                    supported_files.append(f)
            
            if not supported_files:
                self.log_message.emit("在指定文件夹中未找到支持的.docx、.doc或.txt文件")
                return False
                
            self.log_message.emit(f"找到 {len(supported_files)} 个文件待处理")
            
            success_count = 0
            total_files = len(supported_files)
            
            for i, filename in enumerate(supported_files, 1):
                progress = int((i / total_files) * 100)
                self.progress_updated.emit(progress, f"处理文件 {i}/{total_files}: {filename}")
                
                file_path = os.path.join(self.input_path, filename)
                file_ext = os.path.splitext(filename)[1].lower()
                
                if file_ext == '.docx':
                    # 对于docx文件，使用转换器类
                    try:
                        result = self.convert_docx_file(file_path, self.output_folder)
                        if result:
                            success_count += 1
                            
                    except Exception as e:
                        self.log_message.emit(f"处理 {filename} 时出错: {str(e)}")
                
                elif file_ext == '.doc':
                    # 对于doc文件，先转换为docx，再转换                 
                    self.log_message.emit("检测到DOC文件，先转换为DOCX格式...")
                    
                    # 创建临时目录用于存放临时转换的DOCX文件
                    with tempfile.TemporaryDirectory() as temp_dir:
                        docx_path = self.convert_doc_to_docx(file_path, temp_dir)
                        if docx_path:
                            self.log_message.emit("DOC文件转换成功，开始繁简转换...")
                            try:
                                result = self.convert_docx_file(docx_path, self.output_folder)
                                if result:
                                    success_count += 1
                            except Exception as e:
                                self.log_message.emit(f"繁简转换 {filename} 时出错: {str(e)}")
                        else:
                            self.log_message.emit(f"DOC文件 {filename} 转换失败")
                
                elif file_ext == '.txt':
                    if self.convert_txt_file(file_path, self.output_folder):
                        success_count += 1
            
            self.log_message.emit(f"处理完成！成功转换 {success_count}/{total_files} 个文件")
            self.progress_updated.emit(100, "转换完成!")
            return True
        
        else:
            self.log_message.emit("错误：输入的路径既不是有效的文件也不是文件夹")
            return False

    def run(self):
        try:
            success = self.process_files()
            if success:
                self.conversion_finished.emit(True, "转换成功完成")
            else:
                self.conversion_finished.emit(False, "转换过程中出现错误")
        except Exception as e:
            self.conversion_finished.emit(False, f"转换失败: {str(e)}")


class ModernUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.current_theme = "dark"  # 默认主题为暗色
        self.init_ui()
        
    def init_ui(self):
        # 设置窗口属性
        self.setWindowTitle("规范繁体字形转换器 V1.1.2")
        self.setGeometry(100, 100, 900, 750)
        self.setMinimumSize(800, 600)
        
        # 设置窗口图标 - 新增的logo功能
        self.setWindowIcon(QIcon(self.get_logo_path()))
        
        # 应用默认主题
        self.apply_theme(self.current_theme)
        
        # 创建主窗口部件
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        
        # 创建主布局
        main_layout = QVBoxLayout(main_widget)
        main_layout.setSpacing(15)
        
        # 标题区域
        title_label = QLabel("规范繁体字形转换器")
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setObjectName("titleLabel")
        main_layout.addWidget(title_label)
        
        # 创建选项卡
        tab_widget = QTabWidget()
        tab_widget.addTab(self.create_conversion_tab(), "文件转换")
        tab_widget.addTab(self.create_settings_tab(), "设置")
        tab_widget.addTab(self.create_about_tab(), "关于")
        main_layout.addWidget(tab_widget)
        
        # 状态栏
        self.statusBar().showMessage("就绪")
    
    def get_logo_path(self):
        """
        获取logo文件的路径
        程序会按以下顺序查找logo文件：
        1. 与程序同目录下的"logo.ico"
        2. 与程序同目录下的"logo.png"
        3. 程序内部资源（如果没有外部文件，则返回空）
        """
        # 尝试查找logo.ico文件
        logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.ico")
        if os.path.exists(logo_path):
            return logo_path
        
        # 尝试查找logo.png文件
        logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.png")
        if os.path.exists(logo_path):
            return logo_path
        
        # 如果没有找到外部文件，可以创建一个临时的logo
        # 这里我们创建一个简单的程序内建图标作为fallback
        return ""
        
    def create_settings_tab(self):
        """创建设置选项卡"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(20)
        
        # 主题设置区域
        theme_group = QGroupBox("主题设置")
        theme_layout = QVBoxLayout(theme_group)
        theme_layout.setSpacing(15)
        theme_layout.setContentsMargins(15, 20, 15, 15)
        
        # 主题选择说明
        theme_label = QLabel("选择您喜欢的界面主题（仅本次有效）:")
        theme_layout.addWidget(theme_label)
        
        # 主题选择按钮
        theme_button_layout = QHBoxLayout()
        
        # 创建单选按钮
        self.dark_theme_radio = QRadioButton("暗色主题")
        self.light_theme_radio = QRadioButton("浅色主题")
        
        # 设置默认选中暗色主题
        self.dark_theme_radio.setChecked(True)
        
        # 将单选按钮添加到布局
        theme_button_layout.addWidget(self.dark_theme_radio)
        theme_button_layout.addWidget(self.light_theme_radio)
        theme_button_layout.addStretch()
        
        theme_layout.addLayout(theme_button_layout)
        
        # 连接信号
        self.dark_theme_radio.toggled.connect(lambda: self.on_theme_changed("dark"))
        self.light_theme_radio.toggled.connect(lambda: self.on_theme_changed("light"))
        
        layout.addWidget(theme_group)
        
        layout.addStretch()
        return tab
    
    def on_theme_changed(self, theme):
        """主题更改事件处理"""
        if (theme == "dark" and self.dark_theme_radio.isChecked()) or \
           (theme == "light" and self.light_theme_radio.isChecked()):
            self.change_theme(theme)
        
    def apply_theme(self, theme):
        """应用指定主题"""
        if theme == "dark":
            self.apply_dark_theme()
        else:
            self.apply_light_theme()
        
    def apply_dark_theme(self):
        """应用暗色主题"""
        self.current_theme = "dark"
        
        # 设置暗色调色板
        dark_palette = QPalette()
        dark_palette.setColor(QPalette.Window, QColor(43, 43, 43))
        dark_palette.setColor(QPalette.WindowText, Qt.white)
        dark_palette.setColor(QPalette.Base, QColor(30, 30, 30))
        dark_palette.setColor(QPalette.AlternateBase, QColor(43, 43, 43))
        dark_palette.setColor(QPalette.ToolTipBase, Qt.white)
        dark_palette.setColor(QPalette.ToolTipText, Qt.white)
        dark_palette.setColor(QPalette.Text, Qt.white)
        dark_palette.setColor(QPalette.Button, QColor(43, 43, 43))
        dark_palette.setColor(QPalette.ButtonText, Qt.white)
        dark_palette.setColor(QPalette.BrightText, Qt.red)
        dark_palette.setColor(QPalette.Link, QColor(42, 130, 218))
        dark_palette.setColor(QPalette.Highlight, QColor(42, 130, 218))
        dark_palette.setColor(QPalette.HighlightedText, Qt.black)
        QApplication.instance().setPalette(dark_palette)
        
        # 设置暗色样式表
        self.setStyleSheet("""
            QMainWindow {
                background-color: #2b2b2b;
            }
            QWidget {
                background-color: #2b2b2b;
                color: #ffffff;
                font-family: "Microsoft YaHei", sans-serif;
            }
            QPushButton {
                background-color: #375a7f;
                border: none;
                color: white;
                padding: 10px;
                font-size: 14px;
                border-radius: 5px;
                min-width: 100px;
            }
            QPushButton:hover {
                background-color: #4a77a8;
            }
            QPushButton:pressed {
                background-color: #2c4a69;
            }
            QPushButton#startButton {
                background-color: #00bc8c;
                font-weight: bold;
                padding: 12px;
                font-size: 16px;
            }
            QPushButton#startButton:hover {
                background-color: #00e6ac;
            }
            QPushButton#browseButton {
                background-color: #3498db;
            }
            QPushButton#browseButton:hover {
                background-color: #5dade2;
            }
            QLineEdit, QTextEdit {
                background-color: #3c3c3c;
                border: 1px solid #555555;
                color: #ffffff;
                padding: 8px;
                border-radius: 4px;
            }
            QGroupBox {
                border: 1px solid #555555;
                border-radius: 5px;
                margin-top: 1ex;
                font-weight: bold;
                color: #3498db;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #3498db;
            }
            QProgressBar {
                border: 1px solid #555555;
                border-radius: 5px;
                text-align: center;
                height: 20px;
                color: white;
            }
            QProgressBar::chunk {
                background-color: #00bc8c;
                width: 20px;
            }
            QComboBox {
                background-color: #3c3c3c;
                border: 1px solid #555555;
                color: #ffffff;
                padding: 5px;
                border-radius: 4px;
            }
            QComboBox::drop-down {
                border: none;
            }
            QComboBox QAbstractItemView {
                background-color: #3c3c3c;
                color: white;
            }
            QLabel {
                color: #aaaaaa;
            }
            QLabel#titleLabel {
                font-size: 24px;
                font-weight: bold;
                color: #00bc8c;
                margin: 10px;
            }
            QListWidget {
                background-color: #3c3c3c;
                border: 1px solid #555555;
                color: #ffffff;
            }
            QCheckBox:disabled {
                color: #777777;
            }
            QTabWidget::pane {
                border: 1px solid #555555;
                background-color: #2b2b2b;
            }
            QTabBar::tab {
                background-color: #3c3c3c;
                color: #aaaaaa;
                padding: 8px 16px;
                margin-right: 2px;
            }
            QTabBar::tab:selected {
                background-color: #375a7f;
                color: white;
            }
            QTabBar::tab:hover:!selected {
                background-color: #4a77a8;
            }
            QRadioButton {
                color: #aaaaaa;
                padding: 8px;
                font-size: 14px;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
            QRadioButton::indicator:checked {
                background-color: #00bc8c;
                border: 2px solid #555555;
                border-radius: 9px;
            }
            QRadioButton::indicator:unchecked {
                background-color: #3c3c3c;
                border: 2px solid #555555;
                border-radius: 9px;
            }
        """)
        
    def apply_light_theme(self):
        """应用浅色主题"""
        self.current_theme = "light"
        
        # 设置浅色调色板
        light_palette = QPalette()
        light_palette.setColor(QPalette.Window, QColor(240, 240, 240))
        light_palette.setColor(QPalette.WindowText, Qt.black)
        light_palette.setColor(QPalette.Base, Qt.white)
        light_palette.setColor(QPalette.AlternateBase, QColor(245, 245, 245))
        light_palette.setColor(QPalette.ToolTipBase, Qt.white)
        light_palette.setColor(QPalette.ToolTipText, Qt.black)
        light_palette.setColor(QPalette.Text, Qt.black)
        light_palette.setColor(QPalette.Button, QColor(240, 240, 240))
        light_palette.setColor(QPalette.ButtonText, Qt.black)
        light_palette.setColor(QPalette.BrightText, Qt.red)
        light_palette.setColor(QPalette.Link, QColor(0, 120, 215))
        light_palette.setColor(QPalette.Highlight, QColor(0, 120, 215))
        light_palette.setColor(QPalette.HighlightedText, Qt.white)
        QApplication.instance().setPalette(light_palette)
        
        # 设置浅色样式表
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f0f0f0;
            }
            QWidget {
                background-color: #f0f0f0;
                color: #333333;
                font-family: "Microsoft YaHei", sans-serif;
            }
            QPushButton {
                background-color: #4a86e8;
                border: none;
                color: white;
                padding: 10px;
                font-size: 14px;
                border-radius: 5px;
                min-width: 100px;
            }
            QPushButton:hover {
                background-color: #6a9ce8;
            }
            QPushButton:pressed {
                background-color: #3a76c8;
            }
            QPushButton#startButton {
                background-color: #4caf50;
                font-weight: bold;
                padding: 12px;
                font-size: 16px;
            }
            QPushButton#startButton:hover {
                background-color: #66bb6a;
            }
            QPushButton#browseButton {
                background-color: #2196f3;
            }
            QPushButton#browseButton:hover {
                background-color: #42a5f5;
            }
            QLineEdit, QTextEdit {
                background-color: white;
                border: 1px solid #cccccc;
                color: #333333;
                padding: 8px;
                border-radius: 4px;
            }
            QGroupBox {
                border: 1px solid #cccccc;
                border-radius: 5px;
                margin-top: 1ex;
                font-weight: bold;
                color: #555555;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #555555;
            }
            QProgressBar {
                border: 1px solid #cccccc;
                border-radius: 5px;
                text-align: center;
                height: 20px;
                color: #333333;
            }
            QProgressBar::chunk {
                background-color: #4caf50;
                width: 20px;
            }
            QComboBox {
                background-color: white;
                border: 1px solid #cccccc;
                color: #333333;
                padding: 5px;
                border-radius: 4px;
            }
            QComboBox::drop-down {
                border: none;
            }
            QComboBox QAbstractItemView {
                background-color: white;
                color: #333333;
                border: 1px solid #cccccc;
            }
            QLabel {
                color: #555555;
            }
            QLabel#titleLabel {
                font-size: 24px;
                font-weight: bold;
                color: #4caf50;
                margin: 10px;
            }
            QListWidget {
                background-color: white;
                border: 1px solid #cccccc;
                color: #333333;
            }
            QCheckBox:disabled {
                color: #aaaaaa;
            }
            QTabWidget::pane {
                border: 1px solid #cccccc;
                background-color: #f0f0f0;
            }
            QTabBar::tab {
                background-color: #e0e0e0;
                color: #555555;
                padding: 8px 16px;
                margin-right: 2px;
            }
            QTabBar::tab:selected {
                background-color: #4a86e8;
                color: white;
            }
            QTabBar::tab:hover:!selected {
                background-color: #d0d0d0;
            }
            QRadioButton {
                color: #333333;
                padding: 8px;
                font-size: 14px;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
            QRadioButton::indicator:checked {
                background-color: #4caf50;
                border: 2px solid #cccccc;
                border-radius: 9px;
            }
            QRadioButton::indicator:unchecked {
                background-color: white;
                border: 2px solid #cccccc;
                border-radius: 9px;
            }
        """)
        
    def change_theme(self, theme):
        """更改主题"""
        if theme != self.current_theme:
            self.apply_theme(theme)
            self.statusBar().showMessage(f"已切换至{theme}主题")
        
    def create_conversion_tab(self):
        """创建转换选项卡"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setSpacing(20)  # 增加垂直间距
        layout.setContentsMargins(15, 15, 15, 15)  # 增加边距
        
        # 文件选择区域
        file_group = QGroupBox("文件选择")
        file_layout = QVBoxLayout(file_group)
        file_layout.setSpacing(12)  # 增加内部控件间距
        file_layout.setContentsMargins(15, 20, 15, 15)  # 增加内边距，顶部更多
        
        # 输入路径
        input_layout = QHBoxLayout()
        input_layout.setSpacing(10)
        self.input_edit = QLineEdit()
        self.input_edit.setPlaceholderText("请选择要转换的文件或文件夹...")
        input_browse_btn = QPushButton("浏览")
        input_browse_btn.setObjectName("browseButton")
        input_browse_btn.clicked.connect(self.browse_input)
        input_layout.addWidget(QLabel("输入路径:"))
        input_layout.addWidget(self.input_edit)
        input_layout.addWidget(input_browse_btn)
        file_layout.addLayout(input_layout)
        
        # 输出路径
        output_layout = QHBoxLayout()
        output_layout.setSpacing(10)
        self.output_edit = QLineEdit()
        self.output_edit.setPlaceholderText("请选择输出文件夹...")
        output_browse_btn = QPushButton("浏览")
        output_browse_btn.setObjectName("browseButton")
        output_browse_btn.clicked.connect(self.browse_output)
        output_layout.addWidget(QLabel("输出路径:"))
        output_layout.addWidget(self.output_edit)
        output_layout.addWidget(output_browse_btn)
        file_layout.addLayout(output_layout)
        
        layout.addWidget(file_group)
        
        # 转换选项区域
        options_group = QGroupBox("转换选项")
        options_layout = QHBoxLayout(options_group)
        options_layout.setSpacing(30)  # 增加选项之间的水平间距
        options_layout.setContentsMargins(15, 20, 15, 15)  # 增加内边距
        
        # 转换类型选择
        type_layout = QVBoxLayout()
        type_layout.setSpacing(8)
        type_layout.addWidget(QLabel("转换类型:"))
        self.type_combo = QComboBox()
        self.type_combo.addItem("繁体转规范繁体")
        self.type_combo.addItem("繁体旧字形转新字形，但保留异体字不转换")
        self.type_combo.addItem("繁体转规范繁体，但保留文档内原有简体字")
        self.type_combo.addItem("繁体旧字形转新字形，但保留文档内原有简体字和异体字")
        self.type_combo.addItem("繁体转简体")
        type_layout.addWidget(self.type_combo)
        options_layout.addLayout(type_layout)
        
        # 高级选项
        advanced_layout = QVBoxLayout()
        advanced_layout.setSpacing(10)
        
        # 保留格式选项 - 设置为灰色不可用
        self.preserve_format_cb = QCheckBox("保留Word文档的原有格式")
        self.preserve_format_cb.setChecked(True)
        self.preserve_format_cb.setEnabled(False)  # 设置为不可用
        self.preserve_format_cb.setToolTip("此选项已固定启用，不可更改")
        
        # 转换脚注选项 - 设置为灰色不可用
        self.convert_footnotes_cb = QCheckBox("转换Word文档里的脚注和尾注")
        self.convert_footnotes_cb.setChecked(True)
        self.convert_footnotes_cb.setEnabled(False)  # 设置为不可用
        self.convert_footnotes_cb.setToolTip("此选项已固定启用，不可更改")
        
        advanced_layout.addWidget(self.preserve_format_cb)
        advanced_layout.addWidget(self.convert_footnotes_cb)
        options_layout.addLayout(advanced_layout)
        
        layout.addWidget(options_group)
        
        # 控制按钮区域
        control_layout = QHBoxLayout()
        self.start_button = QPushButton("开始转换")
        self.start_button.setObjectName("startButton")
        self.start_button.clicked.connect(self.start_conversion)
        control_layout.addWidget(self.start_button)
        control_layout.addStretch()
        layout.addLayout(control_layout)
        
        # 进度区域
        progress_group = QGroupBox("进度")
        progress_layout = QVBoxLayout(progress_group)
        progress_layout.setSpacing(30)
        progress_layout.setContentsMargins(15, 20, 15, 15)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        self.progress_label = QLabel("准备就绪")
        self.progress_label.setAlignment(Qt.AlignCenter)
        self.log_text = QTextEdit()
        self.log_text.setMaximumHeight(150)
        self.log_text.setReadOnly(True)
        
        progress_layout.addWidget(self.progress_bar)
        progress_layout.addWidget(self.progress_label)
        progress_layout.addWidget(self.log_text)
        
        layout.addWidget(progress_group)
        
        return tab
        
    def create_about_tab(self):
        """创建关于选项卡"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(15, 15, 15, 15)
        
        # 描述区域
        desc_label = QLabel("""
        <h2>规范繁体字形转换器 V1.1.2</h2>
        <p>专业的繁体字形转换工具，助您将繁体旧字形、异体字和港台标准的繁体字形转换为《通用规范汉字表》的规范繁体字形。</p>
        <p><b>主要特性:</b></p>
        <ul>
            <li>支持DOC、DOCX、TXT文档繁体字形转换</li>
            <li>基于《通用规范汉字表》</li>
            <li>转换后保留原文档格式</li>
            <li>支持批量处理文件</li>
            <li>多种转换预设模式适合处理情况复杂的文档</li>
        </ul>
        <p><b>请从以下页面获取本工具最新版本：</p>
        <ul>
              <p>Github：https://github.com/TerryTian-tech/OpenCC-Traditional-Chinese-characters-according-to-Chinese-government-standards
              <p>Gitee：https://gitee.com/terrytian-tech/opencc-traditional-chinese-characters-according-to-chinese-government-standards
        </ul>
        <p><b>本软件遵循Apache-2.0开源协议发布。</p>
        """)
        desc_label.setWordWrap(True)
        desc_label.setAlignment(Qt.AlignLeft)
        layout.addWidget(desc_label)
        
        layout.addStretch()
        return tab
        
    def browse_input(self):
        """浏览输入路径"""
        dialog = QFileDialog()
        dialog.setFileMode(QFileDialog.ExistingFiles)
        dialog.setOption(QFileDialog.ShowDirsOnly, True)
        
        # 允许选择文件或文件夹
        choice = QMessageBox.question(self, "选择类型", "批量转换同一目录下所有文档请选择“Yes”，转换单个文档请选择“No”。",
                                     QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel)
        
        if choice == QMessageBox.Yes:  # 文件夹
            path = QFileDialog.getExistingDirectory(self, "选择输入文件夹")
            if path:
                self.input_edit.setText(path)
        elif choice == QMessageBox.No:  # 文件
            paths, _ = QFileDialog.getOpenFileNames(
                self, "选择文件", "", 
                "文档文件 (*.doc *.docx *.txt);;所有文件 (*)"
            )
            if paths:
                # 如果选择了多个文件，只使用第一个或者让用户选择文件夹
                if len(paths) == 1:
                    self.input_edit.setText(paths[0])
                else:
                    folder = os.path.dirname(paths[0])
                    self.input_edit.setText(folder)
            
    def browse_output(self):
        """浏览输出路径"""
        path = QFileDialog.getExistingDirectory(self, "选择输出文件夹")
        if path:
            self.output_edit.setText(path)
            
    def start_conversion(self):
        """开始转换"""
        input_path = self.input_edit.text()
        output_path = self.output_edit.text()
        
        if not input_path or not output_path:
            QMessageBox.warning(self, "警告", "请输入完整的路径信息")
            return
            
        if not os.path.exists(input_path):
            QMessageBox.critical(self, "错误", "输入路径不存在")
            return
            
        # 获取转换类型
        conversion_types = {
            "繁体转规范繁体": "t2gov",
            "繁体旧字形转新字形，但保留异体字不转换": "t2new",
            "繁体转规范繁体，但保留文档内原有简体字": "t2gov_keep_simp",
            "繁体旧字形转新字形，但保留文档内原有简体字和异体字": "t2new_keep_simp",
            "繁体转简体": "t2s"
        }
        conversion_type = conversion_types[self.type_combo.currentText()]
        
        # 启动转换线程
        self.start_button.setEnabled(False)
        self.progress_bar.setValue(0)
        self.log_text.clear()
        
        # 注意：由于复选框已禁用，我们总是传递True值
        self.worker = ConversionWorker(
            input_path, 
            output_path,
            conversion_type,
            True,  # preserve_format 固定为True
            True   # convert_footnotes 固定为True
        )
        self.worker.progress_updated.connect(self.update_progress)
        self.worker.conversion_finished.connect(self.conversion_finished)
        self.worker.log_message.connect(self.append_log)
        self.worker.start()
        
    def update_progress(self, value, message):
        """更新进度"""
        self.progress_bar.setValue(value)
        self.progress_label.setText(message)
        self.append_log(f"[{value}%] {message}")
        
    def append_log(self, message):
        """添加日志消息"""
        self.log_text.append(message)
        self.log_text.verticalScrollBar().setValue(
            self.log_text.verticalScrollBar().maximum()
        )
        
    def conversion_finished(self, success, message):
        """转换完成"""
        self.start_button.setEnabled(True)
        
        if success:
            QMessageBox.information(self, "成功", message)
            self.statusBar().showMessage("转换完成")
        else:
            QMessageBox.critical(self, "错误", message)
            self.statusBar().showMessage("转换失败")

def main():
    app = QApplication(sys.argv)
    app.setStyle(QStyleFactory.create("Fusion"))  # 使用现代样式
    
    # 设置应用程序图标（会显示在任务栏）
    # 注意：Windows上可能还需要单独的.ico文件才能正确显示任务栏图标
    app_icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.ico")
    if os.path.exists(app_icon_path):
        app.setWindowIcon(QIcon(app_icon_path))
    else:
        # 如果没有找到图标文件，可以尝试png格式
        app_icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.png")
        if os.path.exists(app_icon_path):
            app.setWindowIcon(QIcon(app_icon_path))
    
    window = ModernUI()
    window.show()
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()